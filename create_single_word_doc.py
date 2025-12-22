#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import os
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def add_formatted_content(doc, content):
    """Add formatted content to the document"""
    lines = content.split('\n')

    for line in lines:
        line_stripped = line.strip()

        if not line_stripped:
            continue

        # Check for phase headers (with emoji circles)
        if line_stripped.startswith('🔵 PHASE') or line_stripped.startswith('🟢 PHASE') or line_stripped.startswith('🟡 PHASE'):
            # Main phase header
            heading = doc.add_heading(line_stripped, level=1)
            heading.runs[0].font.color.rgb = RGBColor(0, 51, 153)
            heading.runs[0].font.size = Pt(18)
            continue

        # Check for step numbers (emoji numbers)
        if re.match(r'^[0-9]️⃣', line_stripped):
            # Step header
            heading = doc.add_heading(line_stripped, level=2)
            heading.runs[0].font.color.rgb = RGBColor(0, 102, 204)
            heading.runs[0].font.size = Pt(16)
            continue

        # Check for section headers
        if any(line_stripped.startswith(prefix) for prefix in [
            'Ne Yapar:', 'Detay:', 'Gerçek Kullanım:', 'SQL Query:',
            'Dönen Data', 'Her Connection', 'Sonuç:', 'Bu Flags',
            'Oluşturulan Klasörler:', 'Neden Gerekli:', 'Kod Mantığı:',
            'Örnek:', 'Performance:', 'Önemli:', 'Server Bilgisi',
            'Bağlantı Süreci:', 'Timeout:', 'Hata Durumları:',
            'Bu durumda:', 'Bu Ne İşe Yarar', 'Performance Kazancı:',
            'Örnekle:', 'SFTP Command:', 'Recursive Scan:',
            'Filter Mantığı:', 'Ek Filtreler:', 'Delta Transfer:',
            'SFTP GET Command:', 'Transfer Detayları:', 'Retry Mechanism:',
            'Max retry:', 'Real-world Files:', 'Bulk Insert', 'Transaction:',
            'Neden Bulk Insert:', 'Bu Tablo Ne İşe Yarar:', 'Table Size:',
            'Calculation:', 'SQL Update:', 'Önce:', 'Sonra:',
            'Bir Sonraki Çalışma', 'Bu CRITICAL', 'Cleanup Operations:',
            'Log Summary:', 'Output:', 'Finally Block:', 'Neden Finally:',
            'Synchronization Mechanism:', 'Handler-', 'Engine bekliyor:',
            'Async Notification:', 'Timeline:', 'Code:', 'Thread Pool:',
            'Decompression:', 'Expansion Rate:', 'Java Library:',
            'Error Handling:', 'XSD Schema', 'Validation Process:',
            'Error Count:', 'Generated Statistics:', 'SQL Insert:',
            'Bu Data Ne İşe Yarar:', 'Grafana Dashboard:', 'Operations:',
            'Integration Options:', '🎯 Özet'
        ]):
            # Subsection header
            para = doc.add_paragraph()
            run = para.add_run(line_stripped)
            run.bold = True
            run.font.color.rgb = RGBColor(51, 51, 51)
            run.font.size = Pt(12)
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = Pt(3)
            continue

        # Regular content
        para = doc.add_paragraph()
        run = para.add_run(line_stripped)
        run.font.size = Pt(11)
        para.paragraph_format.line_spacing = 1.15
        para.paragraph_format.space_after = Pt(3)

def main():
    # Read the content file
    with open('akis_ozet.text', 'r', encoding='utf-8') as f:
        content = f.read()

    # Create document
    doc = Document()

    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Add main title
    title = doc.add_heading('Transfer Flow - Tüm Adımlar', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.size = Pt(24)
    title.runs[0].font.color.rgb = RGBColor(0, 51, 153)

    # Add subtitle
    subtitle = doc.add_paragraph('Transfer Workflow Engine: Detaylı Süreç Dokümantasyonu')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].italic = True
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.color.rgb = RGBColor(102, 102, 102)

    # Add a page break after title
    doc.add_page_break()

    # Add all content
    add_formatted_content(doc, content)

    # Save document
    filename = 'Transfer_Flow_Tum_Adimlar.docx'
    doc.save(filename)
    print(f'\n✓ Tek Word dosyası oluşturuldu: {filename}\n')

    # Get file size
    size_mb = os.path.getsize(filename) / (1024 * 1024)
    print(f'  Dosya boyutu: {size_mb:.2f} MB')
    print(f'  Tüm adımlar tek bir dokümanda!\n')

if __name__ == '__main__':
    main()
