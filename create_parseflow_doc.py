from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_heading_with_color(doc, text, level, color_rgb):
    """Renkli başlık ekle"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(*color_rgb)
    return heading

def add_code_block(doc, code_text):
    """Kod bloğu ekle"""
    paragraph = doc.add_paragraph()
    paragraph.style = 'Normal'
    run = paragraph.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

    # Gri arka plan
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), 'F0F0F0')
    paragraph._element.get_or_add_pPr().append(shading_elm)

    return paragraph

def add_table_row(table, cells_data, is_header=False):
    """Tabloya satır ekle"""
    row = table.add_row()
    for i, cell_text in enumerate(cells_data):
        cell = row.cells[i]
        cell.text = cell_text
        if is_header:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
            # Header arka plan rengi
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), '4472C4')
            cell._element.get_or_add_tcPr().append(shading_elm)

# Dokuman oluştur
doc = Document()

# Başlık
title = doc.add_heading('ParseFlow Activity Diagram', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = RGBColor(0, 51, 102)

# Alt başlık
subtitle = doc.add_paragraph('Adım Adım Detaylı Akış Açıklaması')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(14)
subtitle.runs[0].font.color.rgb = RGBColor(68, 114, 196)
subtitle.runs[0].italic = True

doc.add_paragraph()

# İçindekiler
add_heading_with_color(doc, '📋 İçindekiler', 1, (0, 51, 102))
toc_items = [
    '1. Başlangıç (Adım 1-2)',
    '2. Metadata Hazırlık (Adım 3-5)',
    '3. Ana Parse Fazı (Adım 6)',
    '4. Post-Parse Phase (Adım 7-8)',
    '5. Content Date Discovery (Adım 9)',
    '6. Data Loading Phase (Adım 10-14)',
    '7. Post-Loading Operations (Adım 15-17)',
    '8. Performans Metrikleri',
    '9. Özet Tablo'
]
for item in toc_items:
    p = doc.add_paragraph(item, style='List Number')
    p.paragraph_format.left_indent = Inches(0.5)

doc.add_page_break()

# ===========================================
# BAŞLANGIÇ
# ===========================================
add_heading_with_color(doc, '🎬 BAŞLANGIÇ', 1, (0, 102, 204))

# Adım 1
add_heading_with_color(doc, 'Adım 1: startEngine(ParseEngineRecord)', 2, (68, 114, 196))

doc.add_paragraph('Ne yapar?', style='Heading 3')
para = doc.add_paragraph()
para.add_run('Parse engine başlatılır').bold = True
doc.add_paragraph('• Transfer modülü veya Scheduler tarafından tetiklenir')
doc.add_paragraph('• flowId, path\'ler ve konfigürasyon bilgileri içerir')

doc.add_paragraph('Girdi:', style='Heading 3')
add_code_block(doc, '''ParseEngineRecord {
    flowId: 1,
    flowCode: "PARSE_HW_ENB_PM",
    threadCount: 8,
    rawPath: "/data/raw/",
    resultPath: "/data/result/"
}''')

doc.add_paragraph()

# Adım 2
add_heading_with_color(doc, 'Adım 2: preparePaths()', 2, (68, 114, 196))

doc.add_paragraph('Ne yapar?', style='Heading 3')
doc.add_paragraph('Gerekli klasörleri oluşturur')

doc.add_paragraph('Oluşturulan klasörler:', style='Heading 3')
add_code_block(doc, '''/data/raw/          ← XML dosyaları buraya gelir
/data/result/       ← Parse edilen CSV'ler buraya yazılır
/data/error/        ← Hatalı parse'lar buraya''')

doc.add_page_break()

# ===========================================
# METADATA HAZIRLIK
# ===========================================
add_heading_with_color(doc, '📋 METADATA HAZIRLIK', 1, (0, 102, 204))

# Adım 3
add_heading_with_color(doc, 'Adım 3: isActiveFetchTables? (Opsiyonel)', 2, (68, 114, 196))

para = doc.add_paragraph()
para.add_run('Koşul: ').bold = True
para.add_run('isActiveFetchTables = true ise')

doc.add_paragraph('Ne yapar?', style='Heading 3')
doc.add_paragraph('• Database şemasından metadata tablolarını oluşturur')
doc.add_paragraph('• Yeni tablolar varsa t_all_table, t_all_column tablolarına ekler')

doc.add_paragraph('Örnek:', style='Heading 3')
add_code_block(doc, """INSERT INTO t_all_table (table_name, table_description)
VALUES ('t_pm_cell_huawei', 'Huawei eNodeB cell PM data');""")

doc.add_paragraph()

# Adım 4
add_heading_with_color(doc, 'Adım 4: getTables()', 2, (68, 114, 196))

doc.add_paragraph('Ne yapar?', style='Heading 3')
doc.add_paragraph('Parse edilecek tabloların metadata\'sını yükler')

doc.add_paragraph('Database Query:', style='Heading 3')
add_code_block(doc, '''SELECT * FROM t_parse_table
WHERE flow_id = 1 AND is_active = true;

SELECT * FROM t_parse_column
WHERE flow_id = 1 AND is_active = true;''')

doc.add_paragraph('Sonuç:', style='Heading 3')
add_code_block(doc, '''Map<String, ParseMapRecord> tables = {
    "t_pm_cell_huawei": {
        tableName: "t_pm_cell_huawei",
        objectKey: "Cell",
        columns: [
            {name: "rsrp", xmlPath: "measValue/r[0]"},
            {name: "rsrq", xmlPath: "measValue/r[1]"}
        ]
    }
}''')

doc.add_paragraph()

# Adım 5
add_heading_with_color(doc, 'Adım 5: isActivePreParse? (Opsiyonel)', 2, (68, 114, 196))

para = doc.add_paragraph()
para.add_run('Koşul: ').bold = True
para.add_run('isActivePreParse = true ise')

doc.add_paragraph('Ne yapar?', style='Heading 3')
doc.add_paragraph('• Vendor-specific ön işlemler')
doc.add_paragraph('• Örnek: Dosya isimlendirme, validasyon')

doc.add_page_break()

# ===========================================
# ANA PARSE FAZI
# ===========================================
add_heading_with_color(doc, '🔥 ANA PARSE FAZI (Main Parsing Phase)', 1, (204, 0, 0))

add_heading_with_color(doc, 'Adım 6: isActiveOnParse?', 2, (68, 114, 196))

para = doc.add_paragraph()
para.add_run('Koşul: ').bold = True
para.add_run('isActiveOnParse = true ise (genelde her zaman aktif)')

para = doc.add_paragraph()
para.add_run('Bu en önemli faz - 7 alt adımdan oluşur').bold = True

doc.add_paragraph()

# Adım 6.1
add_heading_with_color(doc, 'Adım 6.1: Get Network Nodes from DB', 3, (68, 114, 196))

doc.add_paragraph('Ne yapar?', style='Heading 4')
doc.add_paragraph('Aktif network node\'larını database\'den çeker')

doc.add_paragraph('Database Query:', style='Heading 4')
add_code_block(doc, '''SELECT node_id, node_name
FROM t_network_node
WHERE branch_id = 1 AND is_active = true;''')

doc.add_paragraph('Sonuç:', style='Heading 4')
add_code_block(doc, '''Map<String, Long> networkNodes = {
    "eNodeB_TR_IST_001": 12345,
    "eNodeB_TR_ANK_001": 12346,
    "gNodeB_TR_IST_5G_001": 12347
}''')

para = doc.add_paragraph()
para.add_run('Neden gerekli? ').bold = True
para.add_run('XML\'deki node name\'i database node ID\'sine map etmek için')

doc.add_paragraph()

# Adım 6.2
add_heading_with_color(doc, 'Adım 6.2: Read XML Files from /raw/', 3, (68, 114, 196))

doc.add_paragraph('Ne yapar?', style='Heading 4')
doc.add_paragraph('/raw/ klasöründeki tüm XML dosyalarını listeler')

doc.add_paragraph('Örnek dosyalar:', style='Heading 4')
add_code_block(doc, '''/raw/20260113_eNodeB_PM_001.xml
/raw/20260113_eNodeB_PM_002.xml
/raw/20260113_gNodeB_PM_001.xml
...
Toplam: 150 dosya''')

doc.add_paragraph('Dosya pattern\'leri:', style='Heading 4')
doc.add_paragraph('• *_eNodeB_*.xml → 4G PM (Performance Management)')
doc.add_paragraph('• *_gNodeB_*.xml → 5G PM')
doc.add_paragraph('• *_RNC_*.xml → 3G PM')
doc.add_paragraph('• *_BSC_*.xml → 2G PM')

doc.add_paragraph()

# Adım 6.3
add_heading_with_color(doc, 'Adım 6.3: Create Thread Pool', 3, (68, 114, 196))

doc.add_paragraph('Ne yapar?', style='Heading 4')
doc.add_paragraph('Parallel processing için thread pool oluşturur')

doc.add_paragraph('Kod:', style='Heading 4')
add_code_block(doc, 'ExecutorService executor = Executors.newFixedThreadPool(8);')

doc.add_paragraph('Thread sayısı:', style='Heading 4')
doc.add_paragraph('• Default: 8 thread')
doc.add_paragraph('• Konfigürasyondan alınır (t_parse_engine.on_parse_thread_count)')

doc.add_page_break()

# Adım 6.4
add_heading_with_color(doc, 'Adım 6.4: LOOP - Create & Submit Handlers', 3, (68, 114, 196))

doc.add_paragraph('Ne yapar?', style='Heading 4')
doc.add_paragraph('• Her XML dosyası için parse handler oluşturur ve thread pool\'a gönderir')
para = doc.add_paragraph()
para.add_run('• SEQUENTİAL (sıralı)').bold = True
para.add_run(' - bir dosya ardından diğeri')

doc.add_paragraph('Pseudo kod:', style='Heading 4')
add_code_block(doc, '''for (File xmlFile : xmlFiles) {
    // 1. Parser türünü belirle
    if (xmlFile.contains("eNodeB_PM")) {
        handler = new HwEnbPmXmlParseHandler(...);
    } else if (xmlFile.contains("gNodeB_PM")) {
        handler = new HwGnbPmXmlParseHandler(...);
    }

    // 2. Handler'a node map'i ver
    handler.setNetworkNodeIds(networkNodes);

    // 3. Thread pool'a gönder (NON-BLOCKING)
    executor.submit(handler);
}''')

para = doc.add_paragraph()
para.add_run('Önemli: ').bold = True
para.add_run('Handler\'lar oluşturulurken ')
para.add_run('sıralı').bold = True
para.add_run(', ama çalışırken ')
para.add_run('paralel').bold = True
para.add_run('!')

doc.add_paragraph()

# Adım 6.5
add_heading_with_color(doc, 'Adım 6.5: ⚡ PARALLELIZATION POINT (Fork)', 3, (204, 0, 0))

doc.add_paragraph('Ne yapar?', style='Heading 4')
para = doc.add_paragraph('• Tüm handler\'lar artık ')
para.add_run('PARALEL').bold = True
para.add_run(' olarak çalışmaya başlar')
doc.add_paragraph('• 8 thread aynı anda 8 XML dosyasını parse eder')

doc.add_paragraph('Fork = Parallel Execution:', style='Heading 4')
add_code_block(doc, '''Thread 1 → XML file 1
Thread 2 → XML file 2  } Aynı anda
Thread 3 → XML file 3  } çalışıyor
...
Thread 8 → XML file 8

Bir thread işini bitirince sıradaki dosyayı alır.''')

doc.add_paragraph()

# Adım 6.6
add_heading_with_color(doc, 'Adım 6.6: Handler Thread İşlemleri', 3, (68, 114, 196))

doc.add_paragraph('Her handler şu adımları yapar:', style='Heading 4')

# 6.6.1
doc.add_paragraph('6.6.1: preHandler()', style='Heading 5')
doc.add_paragraph('Dosya adından metadata çıkarır')
add_code_block(doc, '''Dosya: 20260113_1530_eNodeB_TR_IST_001.xml

Çıkarılan:
- fragmentDate: 2026-01-13 15:30
- nodeName: eNodeB_TR_IST_001
- nodeId: 12345 (networkNodes map'inden)''')

doc.add_paragraph()

# 6.6.2
doc.add_paragraph('6.6.2: Open XML with SAX Parser', style='Heading 5')
doc.add_paragraph('SAX parser ile XML dosyasını açar (memory-efficient)')
add_code_block(doc, '''<measInfo>            ← startElement("measInfo")
  <measValue>         ← startElement("measValue")
    75 82 1024000     ← characters("75 82 1024000")
  </measValue>        ← endElement("measValue")
</measInfo>           ← endElement("measInfo")''')

doc.add_paragraph()

# 6.6.3
doc.add_paragraph('6.6.3: Parse XML Elements (Nested Loop)', style='Heading 5')

doc.add_paragraph('Dış döngü: measInfo sections')
add_code_block(doc, '''repeat {
    read measInfo section;
    // İç döngü: measValue records
}''')

doc.add_paragraph('İç döngü: measValue records')
add_code_block(doc, '''repeat {
    Extract metrics (RSRP, RSRQ, Throughput);
    Map to table columns;
    Write to CSV buffer;

    if (autoCounter enabled) {
        Collect counter definitions;
    }
} while (more measValues);''')

doc.add_paragraph('Örnek XML → CSV dönüşümü:')

doc.add_paragraph('XML:')
add_code_block(doc, '''<measInfo measInfoId="Cell">
  <measTypes>RSRP RSRQ Throughput_DL</measTypes>
  <measValue measObjLdn="eNodeB=1,Cell=1">
    <r>75 82 1024000</r>
  </measValue>
</measInfo>''')

doc.add_paragraph('CSV (t_pm_cell_huawei-20260113.csv):')
add_code_block(doc, '12345,001-01,1,1,75,82,1024000,2026-01-13 15:30:00')

doc.add_page_break()

# 6.6.4
doc.add_paragraph('6.6.4: Auto Counter (Opsiyonel)', style='Heading 5')
doc.add_paragraph('Koşul: autoCounter enabled = true ise')
doc.add_paragraph('XML\'de bulunan yeni metric\'leri collect eder')

add_code_block(doc, '''autoCounterDefine.collect(new CounterDefineRecord(
    nodeGroupType: "eNodeB",
    counterGroupType: "Cell",
    counterKey: "L.Cell.RSRP.Mean",
    counterName: "Average RSRP"
));''')

doc.add_paragraph()

# 6.6.5
doc.add_paragraph('6.6.5: postHandler()', style='Heading 5')
doc.add_paragraph('• Resource\'ları temizler')
doc.add_paragraph('• File handle\'ları kapatır')

doc.add_paragraph()

# Adım 6.7
add_heading_with_color(doc, 'Adım 6.7: ⚡ SYNCHRONIZATION POINT (Join)', 3, (204, 0, 0))

doc.add_paragraph('Ne yapar?', style='Heading 4')
doc.add_paragraph('Tüm handler thread\'lerinin bitmesini bekler')

add_code_block(doc, '''executor.shutdown();
executor.awaitTermination(Long.MAX_VALUE, TimeUnit.MILLISECONDS);''')

doc.add_paragraph()

# Adım 6.8
add_heading_with_color(doc, 'Adım 6.8: shutdownExecutorService()', 3, (68, 114, 196))
doc.add_paragraph('Thread pool\'u kapatır')

doc.add_paragraph()

# Adım 6.9
add_heading_with_color(doc, 'Adım 6.9: writer.closeAllStreams()', 3, (68, 114, 196))
doc.add_paragraph('• Tüm CSV buffer\'larını flush eder')
doc.add_paragraph('• File writer\'ları kapatır')

doc.add_paragraph('Sonuç:')
add_code_block(doc, '''/data/result/t_pm_cell_huawei-20260113.csv      (12,500 satır)
/data/result/t_pm_sector_huawei-20260113.csv    (8,300 satır)
/data/result/t_pm_enodeb_huawei-20260113.csv    (450 satır)
...
Toplam: ~50 CSV dosyası''')

doc.add_page_break()

# ===========================================
# POST-PARSE PHASE
# ===========================================
add_heading_with_color(doc, '📅 POST-PARSE PHASE', 1, (0, 102, 204))

# Adım 7
add_heading_with_color(doc, 'Adım 7: isActivePostParse? (Opsiyonel)', 2, (68, 114, 196))
para = doc.add_paragraph()
para.add_run('Koşul: ').bold = True
para.add_run('isActivePostParse = true ise')
doc.add_paragraph('• Vendor-specific post-processing')
doc.add_paragraph('• Aggregation\'lar')
doc.add_paragraph('• Validation\'lar')

doc.add_paragraph()

# Adım 8
add_heading_with_color(doc, 'Adım 8: isActiveAutoCounter? (Opsiyonel)', 2, (68, 114, 196))
para = doc.add_paragraph()
para.add_run('Koşul: ').bold = True
para.add_run('isActiveAutoCounter = true ise')

doc.add_paragraph('Ne yapar?')
doc.add_paragraph('Keşfedilen counter tanımlarını database\'e kaydeder')

add_code_block(doc, '''INSERT INTO t_all_counter
(flow_id, node_group_type, counter_key, counter_name)
SELECT ... FROM temp_counters
WHERE NOT EXISTS (already in t_all_counter);''')

para = doc.add_paragraph('Sonuç: ')
para.add_run('~1000 yeni counter keşfedildi ve kaydedildi').bold = True

doc.add_page_break()

# ===========================================
# CONTENT DATE DISCOVERY
# ===========================================
add_heading_with_color(doc, '📊 CONTENT DATE DISCOVERY', 1, (0, 102, 204))

# Adım 9
add_heading_with_color(doc, 'Adım 9: isActiveDiscoverContentDate? (Opsiyonel)', 2, (68, 114, 196))
para = doc.add_paragraph()
para.add_run('Koşul: ').bold = True
para.add_run('isActiveDiscoverContentDate = true ise')

doc.add_paragraph('Ne yapar?')
para = doc.add_paragraph('CSV dosyalarındaki tarih range\'lerini analiz eder ')
para.add_run('(PARALLEL)').bold = True

doc.add_paragraph('Fork = Parallel Analysis:')
add_code_block(doc, '''Thread 1 → Analyze CSV file 1 (min/max date)
Thread 2 → Analyze CSV file 2
Thread 3 → Analyze CSV file 3
...''')

doc.add_paragraph('Örnek sonuç:')
add_code_block(doc, '''t_pm_cell_huawei-20260113.csv
  Min date: 2026-01-13 00:00:00
  Max date: 2026-01-13 23:45:00''')

doc.add_paragraph('Database Write:')
add_code_block(doc, '''INSERT INTO t_content_date_result
(flow_id, file_name, fragment_date, min_date, max_date)
VALUES (1, 't_pm_cell_huawei-20260113.csv', '2026-01-13', ...);''')

doc.add_page_break()

# ===========================================
# DATA LOADING PHASE
# ===========================================
add_heading_with_color(doc, '💾 DATA LOADING PHASE', 1, (0, 102, 204))

# Adım 10
add_heading_with_color(doc, 'Adım 10: isActiveCleanDuplicateBefore? (Opsiyonel)', 2, (68, 114, 196))
doc.add_paragraph('CSV\'lerdeki duplicate kayıtları temizler')

doc.add_paragraph()

# Adım 11
add_heading_with_color(doc, 'Adım 11: Read CSV Files from /result/', 2, (68, 114, 196))
doc.add_paragraph('Parse edilmiş CSV dosyalarını listeler')

doc.add_paragraph()

# Adım 12
add_heading_with_color(doc, 'Adım 12: PARALLEL - Load to Database (Fork)', 2, (204, 0, 0))

doc.add_paragraph('Ne yapar?')
para = doc.add_paragraph('Her CSV dosyası ')
para.add_run('parallel').bold = True
para.add_run(' olarak database\'e yüklenir')

doc.add_paragraph('Thread Pool: 8 thread (yeni pool)')

doc.add_paragraph('Fork = Parallel Loading:')
add_code_block(doc, '''Thread 1 → LoaderFactory.load(csv1) → Database
Thread 2 → LoaderFactory.load(csv2) → Database
Thread 3 → LoaderFactory.load(csv3) → Database
...''')

doc.add_paragraph('Loader metodları:')
doc.add_paragraph('• PostgreSQL: COPY table FROM csv')
doc.add_paragraph('• Oracle: SQLLDR')
doc.add_paragraph('• MSSQL: BULK INSERT')

doc.add_paragraph('Örnek:')
add_code_block(doc, '''COPY t_pm_cell_huawei FROM '/data/result/t_pm_cell_huawei-20260113.csv'
WITH (FORMAT CSV, DELIMITER ',');''')

doc.add_paragraph('Database Write:')
add_code_block(doc, '''INSERT INTO t_loader_result
(flow_id, file_name, table_name, row_count, status, duration_ms)
VALUES (1, 't_pm_cell_huawei-20260113.csv', 't_pm_cell_huawei',
        12500, 'SUCCESS', 1850);''')

doc.add_paragraph()

# Adım 13
add_heading_with_color(doc, 'Adım 13: shutdownExecutorService()', 2, (68, 114, 196))
doc.add_paragraph('• Loader thread pool\'unu kapatır')
doc.add_paragraph('• Tüm load\'ların bitmesini bekler')

doc.add_paragraph()

# Adım 14
add_heading_with_color(doc, 'Adım 14: isActiveCleanDuplicateAfter? (Opsiyonel)', 2, (68, 114, 196))
doc.add_paragraph('Database\'deki duplicate kayıtları temizler')

add_code_block(doc, '''DELETE FROM t_pm_cell_huawei a
WHERE a.ctid < (
    SELECT max(b.ctid)
    FROM t_pm_cell_huawei b
    WHERE a.node_id = b.node_id AND a.fragment_date = b.fragment_date
);''')

doc.add_page_break()

# ===========================================
# POST-LOADING OPERATIONS
# ===========================================
add_heading_with_color(doc, '🔄 POST-LOADING OPERATIONS', 1, (0, 102, 204))

# Adım 15
add_heading_with_color(doc, 'Adım 15: isActiveCallProcedure? (Opsiyonel)', 2, (68, 114, 196))
doc.add_paragraph('Stored procedure\'leri çalıştırır')
add_code_block(doc, "CALL sp_process_pm_data('2026-01-13');")

doc.add_paragraph()

# Adım 16
add_heading_with_color(doc, 'Adım 16: isActiveCallAggregate? (Opsiyonel)', 2, (68, 114, 196))
doc.add_paragraph('Aggregation query\'leri çalıştırır')

add_code_block(doc, '''-- Hourly averages
INSERT INTO t_pm_cell_huawei_hourly
SELECT node_id, date_trunc('hour', fragment_date),
       AVG(rsrp), AVG(rsrq), SUM(throughput_dl)
FROM t_pm_cell_huawei
WHERE fragment_date = '2026-01-13'
GROUP BY node_id, date_trunc('hour', fragment_date);''')

doc.add_paragraph()

# Adım 17
add_heading_with_color(doc, 'Adım 17: isActiveCallExport? (Opsiyonel)', 2, (68, 114, 196))
doc.add_paragraph('İşlenmiş verileri external sistemlere export eder')

doc.add_paragraph()
doc.add_paragraph()

# Bitiş
para = doc.add_paragraph()
para.add_run('🏁 Parse flow tamamlandı!').bold = True
para.runs[0].font.size = Pt(16)
para.runs[0].font.color.rgb = RGBColor(0, 128, 0)
para.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# ===========================================
# PERFORMANS METRİKLERİ
# ===========================================
add_heading_with_color(doc, '📈 PERFORMANS METRİKLERİ', 1, (0, 102, 204))

doc.add_paragraph('Sequential (Paralelsiz):', style='Heading 3')
add_code_block(doc, '150 dosya × 2 dakika = 300 dakika = 5 saat')

doc.add_paragraph('Parallel (8 thread):', style='Heading 3')
add_code_block(doc, '300 dakika ÷ 8 thread = 37.5 dakika ≈ 38 dakika')

para = doc.add_paragraph()
para.add_run('Performans artışı: %87 daha hızlı! 🚀').bold = True
para.runs[0].font.size = Pt(14)
para.runs[0].font.color.rgb = RGBColor(0, 128, 0)

doc.add_paragraph()
doc.add_paragraph()

# ===========================================
# ÖZET TABLO
# ===========================================
add_heading_with_color(doc, '🎯 ÖZET TABLO', 1, (0, 102, 204))

# Tablo oluştur
table = doc.add_table(rows=1, cols=5)
table.style = 'Light Grid Accent 1'

# Header
add_table_row(table, ['Faz', 'Adım Sayısı', 'Parallel', 'Süre', 'Açıklama'], is_header=True)

# Satırlar
rows_data = [
    ['Başlangıç', '2', '❌', '<1 dk', 'Engine başlatma, path oluşturma'],
    ['Metadata Hazırlık', '3', '❌', '<1 dk', 'Tablo metadata yükleme'],
    ['Ana Parse', '9', '✅', '~35 dk', 'XML parsing (8 thread parallel)'],
    ['Post-Parse', '2', '❌', '<1 dk', 'Auto counter kaydetme'],
    ['Content Date', '1', '✅', '~1 dk', 'Date range discovery'],
    ['Data Loading', '5', '✅', '~5 dk', 'CSV → Database (parallel)'],
    ['Post-Loading', '3', '❌', '~5 dk', 'Procedures, aggregations'],
    ['TOPLAM', '25', '3 faz', '~48 dk', '150 XML dosyası']
]

for row_data in rows_data:
    add_table_row(table, row_data, is_header=False)

doc.add_paragraph()
doc.add_paragraph()

# Footer
para = doc.add_paragraph()
para.add_run('Bu akış sayesinde 150 XML dosyası sadece 48 dakikada parse edilip database\'e yüklenir! 🎉').bold = True
para.runs[0].font.size = Pt(12)
para.runs[0].font.color.rgb = RGBColor(0, 102, 204)
para.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Son sayfa - Bilgi
doc.add_page_break()
add_heading_with_color(doc, '📚 Döküman Bilgileri', 1, (0, 51, 102))

info_table = doc.add_table(rows=6, cols=2)
info_table.style = 'Light List Accent 1'

info_data = [
    ['Döküman Adı', 'ParseFlow Activity Diagram - Detaylı Akış'],
    ['Tarih', '2026-01-13'],
    ['Versiyon', '1.0'],
    ['Kaynak', 'ParseFlow_ActivityDiagram.puml'],
    ['Modül', 'Parser Module - Parse Flow'],
    ['Toplam Adım', '25 (3 parallel faz)']
]

for i, (key, value) in enumerate(info_data):
    row = info_table.rows[i]
    row.cells[0].text = key
    row.cells[1].text = value
    # Key bold yap
    for paragraph in row.cells[0].paragraphs:
        for run in paragraph.runs:
            run.font.bold = True

# Dosyayı kaydet
doc.save('ParseFlow_Activity_Diagram_Detayli_Akis.docx')
print("✅ Word dokümanı oluşturuldu: ParseFlow_Activity_Diagram_Detayli_Akis.docx")
