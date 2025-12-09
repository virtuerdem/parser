#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import os
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def parse_content(content):
    """Parse the content into phases and steps"""
    steps = []

    # Split by step numbers (emoji numbers like 1️⃣, 2️⃣, etc.)
    pattern = r'([0-9]️⃣.*?)(?=[0-9]️⃣|🟢 PHASE|🟡 PHASE|🎯 Özet|$)'
    matches = re.findall(pattern, content, re.DOTALL)

    for match in matches:
        steps.append(match.strip())

    return steps

def create_word_document(step_number, step_content, output_dir):
    """Create a Word document for a single step"""
    doc = Document()

    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Extract title from step content (first line)
    lines = step_content.split('\n')
    title = lines[0] if lines else f"Adım {step_number}"

    # Add title
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add step content
    current_para = None
    in_code_block = False

    for line in lines[1:]:
        line = line.strip()

        if not line:
            continue

        # Check if it's a section header
        if line.startswith('Ne Yapar:') or line.startswith('Detay:') or \
           line.startswith('Gerçek Kullanım:') or line.startswith('SQL Query:') or \
           line.startswith('Dönen Data') or line.startswith('Her Connection') or \
           line.startswith('Sonuç:') or line.startswith('Bu Flags') or \
           line.startswith('Oluşturulan Klasörler:') or line.startswith('Neden Gerekli:') or \
           line.startswith('Kod Mantığı:') or line.startswith('Örnek:') or \
           line.startswith('Performance:') or line.startswith('Önemli:') or \
           line.startswith('Server Bilgisi') or line.startswith('Bağlantı Süreci:') or \
           line.startswith('Timeout:') or line.startswith('Hata Durumları:') or \
           line.startswith('Bu durumda:') or line.startswith('Bu Ne İşe Yarar') or \
           line.startswith('Performance Kazancı:') or line.startswith('Örnekle:') or \
           line.startswith('SFTP Command:') or line.startswith('Dönen Data') or \
           line.startswith('Recursive Scan:') or line.startswith('Filter Mantığı:') or \
           line.startswith('Ek Filtreler:') or line.startswith('Delta Transfer:') or \
           line.startswith('SFTP GET Command:') or line.startswith('Transfer Detayları:') or \
           line.startswith('Retry Mechanism:') or line.startswith('Max retry:') or \
           line.startswith('Real-world Files:') or line.startswith('Bulk Insert') or \
           line.startswith('Transaction:') or line.startswith('Neden Bulk Insert:') or \
           line.startswith('Bu Tablo Ne İşe Yarar:') or line.startswith('Table Size:') or \
           line.startswith('Calculation:') or line.startswith('SQL Update:') or \
           line.startswith('Önce:') or line.startswith('Sonra:') or \
           line.startswith('Bir Sonraki Çalışma') or line.startswith('Bu CRITICAL') or \
           line.startswith('Cleanup Operations:') or line.startswith('Log Summary:') or \
           line.startswith('Output:') or line.startswith('Finally Block:') or \
           line.startswith('Neden Finally:') or line.startswith('Synchronization Mechanism:') or \
           line.startswith('Handler-') or line.startswith('Engine bekliyor:') or \
           line.startswith('Async Notification:') or line.startswith('Timeline:') or \
           line.startswith('Code:') or line.startswith('Thread Pool:') or \
           line.startswith('Decompression:') or line.startswith('Expansion Rate:') or \
           line.startswith('Java Library:') or line.startswith('Error Handling:') or \
           line.startswith('XSD Schema') or line.startswith('Validation Process:') or \
           line.startswith('Neden Gerekli:') or line.startswith('Error Count:') or \
           line.startswith('Generated Statistics:') or line.startswith('SQL Insert:') or \
           line.startswith('Bu Data Ne İşe Yarar:') or line.startswith('Grafana Dashboard:') or \
           line.startswith('Operations:') or line.startswith('Integration Options:'):
            heading = doc.add_heading(line, level=2)
            heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)
            current_para = None
        else:
            # Regular paragraph
            if current_para is None:
                current_para = doc.add_paragraph()
            else:
                current_para = doc.add_paragraph()

            current_para.add_run(line)
            current_para.paragraph_format.line_spacing = 1.15

    # Save document
    filename = f"Adim_{step_number:02d}_{title[:30].replace('/', '_').replace(':', '_').replace('?', '')}.docx"
    filepath = os.path.join(output_dir, filename)
    doc.save(filepath)
    print(f"✓ Created: {filename}")

    return filepath

def main():
    # Read the content file
    with open('akis_ozet.text', 'r', encoding='utf-8') as f:
        content = f.read()

    # Create output directory
    output_dir = 'word_docs'
    os.makedirs(output_dir, exist_ok=True)

    # Parse steps
    steps = parse_content(content)

    print(f"\nFound {len(steps)} steps. Creating Word documents...\n")

    # Create a document for each step
    for i, step in enumerate(steps, 1):
        if step.strip():
            create_word_document(i, step, output_dir)

    print(f"\n✓ All done! Created {len(steps)} Word documents in '{output_dir}/' directory\n")

if __name__ == '__main__':
    main()
