#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def create_parser_db_tables_word():
    doc = Document()

    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Title
    title = doc.add_heading('Parser Modülü - Database Tabloları', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.size = Pt(24)
    title.runs[0].font.color.rgb = RGBColor(0, 51, 153)

    # Subtitle
    subtitle = doc.add_paragraph('Parser Modülünde Kullanılan Tüm Database Tablolarının Detaylı Listesi')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].italic = True
    subtitle.runs[0].font.size = Pt(12)

    doc.add_page_break()

    # Table of Contents
    doc.add_heading('İçindekiler', level=1)
    toc_items = [
        '1. Metadata & Configuration Tables',
        '2. Monitoring & History Tables',
        '3. Network & Reference Tables',
        '4. Data Tables (Parse Output)',
        '5. Tablo İlişkileri',
        '6. Örnek SQL Queryler'
    ]
    for item in toc_items:
        p = doc.add_paragraph(item, style='List Number')
        p.runs[0].font.size = Pt(11)

    doc.add_page_break()

    # Overview
    doc.add_heading('Genel Bakış', level=1)

    overview_text = """Parser modülü 60+ database tablosu kullanır. Bu tablolar 4 ana kategoriye ayrılır:

• Metadata & Configuration Tables (7 tablo) - Parser konfigürasyonu ve metadata
• Monitoring & History Tables (3 tablo) - İzleme, log ve geçmiş
• Network & Reference Tables (4 tablo) - Referans veriler
• Data Tables (50+ tablo) - Parse edilen verinin yazıldığı hedef tablolar"""

    p = doc.add_paragraph(overview_text)
    p.runs[0].font.size = Pt(11)

    # Category 1: Metadata & Configuration
    doc.add_page_break()
    doc.add_heading('1. Metadata & Configuration Tables', level=1)

    # t_parse_engine
    doc.add_heading('t_parse_engine', level=2)
    doc.add_paragraph('Amaç: Parse engine konfigürasyonu', style='Intense Quote')

    p = doc.add_paragraph()
    p.add_run('Önemli Kolonlar:\n').bold = True
    cols = [
        'flow_id - Hangi flow için',
        'is_active_on_parse - Ana parsing aktif mi?',
        'on_parse_thread_count - XML parsing thread sayısı (default: 8)',
        'loader_thread_count - DB load thread sayısı (default: 8)',
        'is_active_auto_counter - Auto counter discovery aktif mi?',
        'is_active_discover_content_date - Date discovery aktif mi?'
    ]
    for col in cols:
        doc.add_paragraph(col, style='List Bullet')

    p = doc.add_paragraph()
    p.add_run('Örnek Data:\n').bold = True
    example = """flow_id: 111
is_active_on_parse: true
on_parse_thread_count: 8
is_active_auto_counter: true"""
    code_para = doc.add_paragraph(example)
    code_para.style = 'No Spacing'
    for run in code_para.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)

    # t_parse_table
    doc.add_heading('t_parse_table', level=2)
    doc.add_paragraph('Amaç: Parse edilecek XML\'lerdeki her bir tablo için metadata', style='Intense Quote')

    p = doc.add_paragraph()
    p.add_run('Önemli Kolonlar:\n').bold = True
    cols = [
        'table_name - Hedef tablo ismi (örn: t_pm_cell_huawei)',
        'object_key - Anahtar (örn: Cell)',
        'node_type - Node tipi (eNodeB, gNodeB, RNC, BSC)',
        'table_type - Tablo tipi (PM, CM, Conf)',
        'network_type - Network tipi (4G, 5G, 3G, 2G)',
        'data_source - Veri kaynağı (Huawei, Ericsson)'
    ]
    for col in cols:
        doc.add_paragraph(col, style='List Bullet')

    # t_parse_column
    doc.add_heading('t_parse_column', level=2)
    doc.add_paragraph('Amaç: Her tablodaki kolonların metadata\'sı', style='Intense Quote')

    p = doc.add_paragraph()
    p.add_run('Önemli Kolonlar:\n').bold = True
    cols = [
        'column_name - Kolon ismi (örn: rsrp, rsrq)',
        'column_type - Veri tipi (integer, varchar, timestamp)',
        'xml_path - XML\'deki path (örn: measResults[0])',
        'column_index - Sıra numarası'
    ]
    for col in cols:
        doc.add_paragraph(col, style='List Bullet')

    # t_all_counter
    doc.add_heading('t_all_counter', level=2)
    doc.add_paragraph('Amaç: Auto-discovery ile bulunan counter/metrik tanımları', style='Intense Quote')

    p = doc.add_paragraph()
    p.add_run('Auto-Discovery Özelliği:\n').bold = True
    p = doc.add_paragraph("""Parser, XML dosyalarından otomatik olarak yeni metrikleri keşfeder ve bu tabloya kaydeder.
Vendor yeni bir metrik eklediğinde, sistem otomatik olarak adapte olur.""")
    p.runs[0].font.size = Pt(11)

    p = doc.add_paragraph()
    p.add_run('Örnek:\n').bold = True
    example = """XML'de yeni metrik:
<measTypes>RSRP RSRQ NewMetric_XYZ</measTypes>

Auto-discovery:
INSERT INTO t_all_counter (counter_key, ...)
VALUES ('NewMetric_XYZ', ...);  ← Otomatik keşfedildi!"""
    code_para = doc.add_paragraph(example)
    for run in code_para.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)

    # Category 2: Monitoring & History
    doc.add_page_break()
    doc.add_heading('2. Monitoring & History Tables', level=1)

    # t_parse_process_history
    doc.add_heading('t_parse_process_history', level=2)
    doc.add_paragraph('Amaç: Her parse çalıştırmasının istatistikleri', style='Intense Quote')

    p = doc.add_paragraph()
    p.add_run('Önemli Kolonlar:\n').bold = True
    cols = [
        'flow_process_code - Unique run ID',
        'total_files - Toplam parse edilen dosya sayısı',
        'success_count - Başarılı dosya sayısı',
        'failure_count - Başarısız dosya sayısı',
        'total_records - Toplam kayıt sayısı',
        'execution_duration_ms - Çalışma süresi (milisaniye)',
        'start_time, end_time - Başlangıç ve bitiş zamanları'
    ]
    for col in cols:
        doc.add_paragraph(col, style='List Bullet')

    p = doc.add_paragraph()
    p.add_run('Örnek Data:\n').bold = True
    example = """flow_process_code: 20240708100000000111
total_files: 150
success_count: 149
failure_count: 1
execution_duration_ms: 2280000 (38 dakika)
total_records: 10000000"""
    code_para = doc.add_paragraph(example)
    for run in code_para.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)

    # t_content_date_result
    doc.add_heading('t_content_date_result', level=2)
    doc.add_paragraph('Amaç: Parse edilen dosyalardaki tarih aralıkları', style='Intense Quote')

    p = doc.add_paragraph("""Parse sonrası CSV dosyalarındaki tarih kolonları analiz edilir ve
her dosya için min/max tarih aralıkları bu tabloya kaydedilir.""")
    p.runs[0].font.size = Pt(11)

    # Category 3: Network & Reference
    doc.add_page_break()
    doc.add_heading('3. Network & Reference Tables', level=1)

    # t_network_node
    doc.add_heading('t_network_node', level=2)
    doc.add_paragraph('Amaç: Aktif network node\'ları (eNodeB, gNodeB, RNC, BSC)', style='Intense Quote')

    p = doc.add_paragraph()
    p.add_run('Kullanım:\n').bold = True
    p = doc.add_paragraph("""Parse sırasında node_name'den node_id'ye mapping yapmak için kullanılır.
XML dosyasındaki node ismi bu tabloda aranır ve node_id bulunur.""")
    p.runs[0].font.size = Pt(11)

    p = doc.add_paragraph()
    p.add_run('Örnek Mapping:\n').bold = True
    example = """eNodeB_001 → node_id: 12345
eNodeB_002 → node_id: 12346
gNodeB_001 → node_id: 12347"""
    code_para = doc.add_paragraph(example)
    for run in code_para.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)

    # Category 4: Data Tables
    doc.add_page_break()
    doc.add_heading('4. Data Tables (Parse Output)', level=1)

    p = doc.add_paragraph("""Parse edilen verilerin gerçekte yazıldığı tablolar.
Bu tablolar vendor ve teknolojiye göre değişir. Parser modülü bu tablolara CSV formatında veri yazar,
sonra bulk load eder.""")
    p.runs[0].font.size = Pt(11)

    # Huawei eNodeB (4G)
    doc.add_heading('Huawei eNodeB (4G) - Performance Management', level=2)

    # t_pm_cell_huawei
    doc.add_heading('t_pm_cell_huawei', level=3)
    doc.add_paragraph('Amaç: Cell-level performans metrikleri', style='Intense Quote')

    p = doc.add_paragraph()
    p.add_run('Tipik Kolonlar:\n').bold = True
    cols = [
        'node_id - FK → t_network_node',
        'fragment_date - Ölçüm zamanı (15 dakikalık period)',
        'plmn, enodeb_id, cell_id - Cell tanımlayıcılar',
        'rsrp - Reference Signal Received Power (dBm)',
        'rsrq - Reference Signal Received Quality (dB)',
        'throughput_dl, throughput_ul - Throughput (Kbps)',
        'prb_usage_dl, prb_usage_ul - PRB kullanımı (%)',
        'active_users - Aktif kullanıcı sayısı',
        'handover_success_rate - Handover başarı oranı (%)'
    ]
    for col in cols:
        doc.add_paragraph(col, style='List Bullet')

    p = doc.add_paragraph()
    p.add_run('Örnek Data:\n').bold = True
    example = """node_id: 12345
fragment_date: 2024-07-08 10:00:00
cell_id: 1
rsrp: -75 dBm
rsrq: -8 dB
throughput_dl: 102400 Kbps
active_users: 120"""
    code_para = doc.add_paragraph(example)
    for run in code_para.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)

    # Huawei gNodeB (5G)
    doc.add_heading('Huawei gNodeB (5G)', level=2)

    doc.add_heading('t_pm_nr_cell_huawei', level=3)
    doc.add_paragraph('Amaç: 5G NR cell performans metrikleri', style='Intense Quote')

    p = doc.add_paragraph()
    p.add_run('5G Spesifik Metrikler:\n').bold = True
    cols = [
        'ssb_rsrp, ssb_rsrq, ssb_sinr - 5G sinyal kalitesi',
        'throughput_dl - 5G yüksek hız (multi-Gbps)',
        'beam_management_metrics - Beam yönetimi',
        'massive_mimo_metrics - Massive MIMO metrikleri'
    ]
    for col in cols:
        doc.add_paragraph(col, style='List Bullet')

    # Table Relationships
    doc.add_page_break()
    doc.add_heading('5. Tablo İlişkileri', level=1)

    p = doc.add_paragraph()
    p.add_run('Parse Flow ve Tablo Kullanımı:\n').bold = True

    flow_text = """1. startEngine()
   READ → t_parse_engine (config)
   READ → t_parse_component (paths)

2. getTables()
   READ → t_parse_table (table mappings)
   READ → t_parse_column (column mappings)

3. onEngine()
   READ → t_network_node (active nodes)

   PARSE 150 XML files
   ↓
   WRITE CSV files:
   - t_pm_cell_huawei-20240708.csv
   - t_pm_sector_huawei-20240708.csv
   - t_cm_cell_huawei-20240708.csv

4. Auto Counter Discovery
   WRITE → t_all_counter

5. Content Date Discovery
   WRITE → t_content_date_result

6. Bulk Load
   BULK INSERT:
   - t_pm_cell_huawei (10M records)
   - t_pm_sector_huawei (5M records)
   - t_cm_cell_huawei (100K records)

7. Save Statistics
   WRITE → t_parse_process_history
   WRITE → t_loader_result"""

    code_para = doc.add_paragraph(flow_text)
    for run in code_para.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)

    # Typical Data Volumes
    doc.add_heading('Tipik Veri Boyutları', level=2)

    # Create table
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Light Grid Accent 1'

    # Header
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Tablo'
    hdr_cells[1].text = 'Kayıt/Ay'
    hdr_cells[2].text = 'Boyut'

    # Data
    data = [
        ('t_pm_cell_huawei', '~10M', '5-10 GB'),
        ('t_pm_sector_huawei', '~5M', '3-5 GB'),
        ('t_cm_cell_huawei', '~100K', '50-100 MB'),
        ('t_all_counter', '~5,000', '5-10 MB'),
        ('t_parse_table', '~50', '<1 MB')
    ]

    for i, (tbl, rec, size) in enumerate(data, 1):
        row = table.rows[i]
        row.cells[0].text = tbl
        row.cells[1].text = rec
        row.cells[2].text = size

    # SQL Examples
    doc.add_page_break()
    doc.add_heading('6. Örnek SQL Queryler', level=1)

    # Query 1
    doc.add_heading('Parse Configuration Sorgulama', level=2)
    query1 = """SELECT pe.*, pc.*
FROM t_parse_engine pe
JOIN t_parse_component pc ON pc.flow_id = pe.flow_id
WHERE pe.flow_id = 111;"""
    code_para = doc.add_paragraph(query1)
    for run in code_para.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)

    # Query 2
    doc.add_heading('Table Metadata Alma', level=2)
    query2 = """SELECT pt.table_name, pt.object_key,
       pc.column_name, pc.column_type
FROM t_parse_table pt
JOIN t_parse_column pc ON pc.parse_table_id = pt.id
WHERE pt.flow_id = 111 AND pt.is_active = true
ORDER BY pt.table_name, pc.column_index;"""
    code_para = doc.add_paragraph(query2)
    for run in code_para.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)

    # Query 3
    doc.add_heading('Parse İstatistikleri (Son 30 Gün)', level=2)
    query3 = """SELECT
    flow_id,
    COUNT(*) as total_runs,
    AVG(total_files) as avg_files,
    AVG(execution_duration_ms / 1000 / 60) as avg_duration_min,
    SUM(total_records) as total_records
FROM t_parse_process_history
WHERE created_time >= CURRENT_DATE - INTERVAL '30 days'
GROUP BY flow_id;"""
    code_para = doc.add_paragraph(query3)
    for run in code_para.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)

    # Query 4
    doc.add_heading('Parse Edilen Veri Sorgulama', level=2)
    query4 = """SELECT
    n.node_name,
    p.fragment_date,
    p.cell_id,
    p.rsrp,
    p.rsrq,
    p.throughput_dl,
    p.active_users
FROM t_pm_cell_huawei p
JOIN t_network_node n ON n.id = p.node_id
WHERE p.fragment_date >= '2024-07-08 00:00:00'
  AND p.fragment_date < '2024-07-09 00:00:00'
  AND n.node_name = 'eNodeB_001'
ORDER BY p.fragment_date;"""
    code_para = doc.add_paragraph(query4)
    for run in code_para.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)

    # Summary
    doc.add_page_break()
    doc.add_heading('Özet', level=1)

    summary_text = """Parser modülü toplam 60+ database tablosu kullanır:

• 7 Metadata & Configuration Tables - Parser konfigürasyonu
• 3 Monitoring & History Tables - İzleme ve istatistikler
• 4 Network & Reference Tables - Referans veriler
• 50+ Data Tables - Parse edilen gerçek veriler

Parser modülü:
✓ XML dosyalarını SAX parser ile parse eder
✓ CSV formatında output üretir
✓ Bulk load ile veritabanına yükler
✓ Yeni metrikleri otomatik keşfeder (auto-discovery)
✓ 8 thread ile parallel çalışır
✓ Tipik performance: 150 XML dosyası → 38 dakika"""

    p = doc.add_paragraph(summary_text)
    p.runs[0].font.size = Pt(11)

    # Save
    filename = 'Parser_Database_Tables.docx'
    doc.save(filename)
    print(f'\n✓ Word dosyası oluşturuldu: {filename}')

    import os
    size_mb = os.path.getsize(filename) / (1024 * 1024)
    print(f'  Dosya boyutu: {size_mb:.2f} MB\n')

if __name__ == '__main__':
    create_parser_db_tables_word()
